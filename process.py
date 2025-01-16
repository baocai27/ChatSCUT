import os
import requests
import base64
from PIL import Image
from docx import Document
from tqdm import tqdm

API_KEY = "RDpeVx8FBuL905G1BRWKv6tH"
SECRET_KEY = "Ery7njEUG3kdXXSk9bKLaZY5LXjaANMq"

def get_access_token():
    """
    使用 AK，SK 生成鉴权签名（Access Token）
    :return: access_token，或是None(如果错误)
    """
    url = "https://aip.baidubce.com/oauth/2.0/token"
    params = {"grant_type": "client_credentials", "client_id": API_KEY, "client_secret": SECRET_KEY}
    response = requests.post(url, params=params)
    return response.json().get("access_token")

def extract_text_from_image(image_path, access_token):
    """
    使用百度OCR接口提取图片中的文字
    :param image_path: 图片路径
    :param access_token: 百度API的access token
    :return: 图片中的文字
    """
    url = "https://aip.baidubce.com/rest/2.0/ocr/v1/accurate_basic?access_token=" + access_token

    with open(image_path, 'rb') as img_file:
        img_data = img_file.read()
    
    img_base64 = base64.b64encode(img_data).decode()

    headers = {'Content-Type': 'application/x-www-form-urlencoded'}
    data = {'image': img_base64}
    response = requests.post(url, headers=headers, data=data)
    result = response.json()

    # 打印返回结果以便调试
    print(result)

    if 'words_result' in result:
        return '\n'.join([item['words'] for item in result['words_result']])
    else:
        return "No text found"

def save_text_to_word(text, output_path):
    """
    将提取的文字保存到Word文档
    :param text: 提取的文字
    :param output_path: Word文档保存路径
    """
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run(text)
    font = run.font
    font.name = '宋体'
    doc.save(output_path)

def main():
    image_folder = "C:/Users/13157/Desktop/images"
    output_folder = "C:/Users/13157/Desktop/doc1"

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    image_files = [f for f in os.listdir(image_folder) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff'))]

    total_images = len(image_files)
    access_token = get_access_token()

    for i, image_file in enumerate(tqdm(image_files, desc="Processing images", unit="image")):
        image_path = os.path.join(image_folder, image_file)
        text = extract_text_from_image(image_path, access_token)
        
        # 打印提取的文字以便调试
        print(f"Extracted text from {image_file}: {text}")

        doc_name = f'{os.path.splitext(image_file)[0]}.docx'
        output_path = os.path.join(output_folder, doc_name)
        
        save_text_to_word(text, output_path)

    print("All images have been processed.")

if __name__ == '__main__':
    main()

